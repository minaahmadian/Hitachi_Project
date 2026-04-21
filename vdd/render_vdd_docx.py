from __future__ import annotations

import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from traceability.pre_isa_report import EVIDENCE_CHAIN_TEXT

from vdd.regulatory_findings_text import format_regulatory_findings_plain


def _templates_dir(repo_root: Path) -> Path:
    return repo_root / "templates"


def ensure_default_vdd_template(repo_root: Path) -> Path:
    """
    Create a minimal docxtpl template (Jinja placeholders) if missing.
    Uses python-docx (already a project dependency).
    """
    templates = _templates_dir(repo_root)
    templates.mkdir(parents=True, exist_ok=True)
    dest = templates / "vdd_pre_isa_template.docx"
    if dest.is_file():
        return dest

    from docx import Document

    doc = Document()
    doc.add_heading("Version Description Document — Pre-ISA Draft", 0)
    doc.add_paragraph(
        "Hitachi / railway software — CENELEC EN 50128–oriented pre-audit bundle (automated draft, not signed)."
    )
    doc.add_heading("Inputs and evidence chain", level=1)
    doc.add_paragraph("{{ evidence_chain_text }}")
    doc.add_heading("1. Release decision", level=1)
    doc.add_paragraph("Verdict: {{ final_decision }}")
    doc.add_paragraph("Lead assessor rationale:")
    doc.add_paragraph("{{ vdd_explanation }}")
    doc.add_heading("2. Pre-ISA consolidated assessment", level=1)
    doc.add_paragraph("Overall gate: {{ pre_isa_overall }}")
    doc.add_paragraph("{{ summary_for_vdd }}")
    doc.add_heading("3. Supporting gate status (deterministic)", level=1)
    doc.add_paragraph("Traceability matcher: {{ matcher_status }}")
    doc.add_paragraph("Regulatory rule engine: {{ regulatory_status }} (derogation_needed={{ regulatory_derogation_needed }})")
    doc.add_paragraph("Communications triage: {{ detective_status }}")
    doc.add_paragraph("Derogation language scan: {{ derogation_scan_overall }}")
    doc.add_heading("4. Per-anomaly verdicts", level=1)
    doc.add_paragraph("{{ verdict_per_anomaly_text }}")
    doc.add_heading("5. Citations (clauses + governance snippets)", level=1)
    doc.add_paragraph("{{ citations_text }}")
    doc.add_paragraph("Document generated at: {{ document_generated_at }}")
    doc.save(str(dest))
    return dest


def build_vdd_template_context(final_state: dict[str, Any]) -> dict[str, Any]:
    pre = final_state.get("pre_isa_report") or {}
    if not isinstance(pre, dict):
        pre = {}

    assessor = final_state.get("assessor_report") or {}
    if not isinstance(assessor, dict):
        assessor = {}

    matcher = final_state.get("matcher_report") or {}
    regulatory = final_state.get("regulatory_report") or {}
    detective = final_state.get("detective_report") or {}
    derogation = final_state.get("derogation_report") or {}

    digest = pre.get("inputs_digest") if isinstance(pre.get("inputs_digest"), dict) else {}

    verdict_lines: list[str] = []
    for row in pre.get("verdict_per_anomaly") or []:
        if not isinstance(row, dict):
            continue
        verdict_lines.append(
            f"• {row.get('anomaly_id', '')} [{row.get('matcher_severity', '')}] "
            f"{row.get('verdict', '')} — {str(row.get('rationale', ''))[:400]}"
        )
    verdict_text = "\n".join(verdict_lines) if verdict_lines else "(none)"

    cite_lines: list[str] = []
    for c in pre.get("citations") or []:
        if not isinstance(c, dict):
            continue
        kind = c.get("kind", "")
        if kind == "regulatory_clause":
            cite_lines.append(
                f"• Clause {c.get('clause_id', '')} ({c.get('title', '')}) score={c.get('score', '')}"
            )
        else:
            cite_lines.append(
                f"• {c.get('pattern_id', '')} [{c.get('strength', '')}] "
                f"{c.get('source', '')}: {str(c.get('snippet', ''))[:180]}"
            )
    citations_text = "\n".join(cite_lines) if cite_lines else "(none)"

    rac_doc = max(120, int(os.getenv("VDD_REGULATORY_RATIONALE_MAX_CHARS", "500") or "500"))
    regulatory_findings_text = format_regulatory_findings_plain(
        regulatory if isinstance(regulatory, dict) else {},
        rationale_max=rac_doc,
    ).strip() or "(no per-rule regulatory findings for this run)"

    ev = str(pre.get("evidence_chain_text", "")).strip() or EVIDENCE_CHAIN_TEXT

    return {
        "final_decision": str(assessor.get("final_decision", "UNKNOWN")),
        "vdd_explanation": str(assessor.get("vdd_explanation", "")).strip() or "(no rationale recorded)",
        "pre_isa_overall": str(pre.get("overall", "UNKNOWN")),
        "evidence_chain_text": ev,
        "summary_for_vdd": str(pre.get("summary_for_vdd", "")).strip() or "(no pre-ISA summary)",
        "matcher_status": str(digest.get("matcher_status", matcher.get("status", ""))),
        "regulatory_status": str(digest.get("regulatory_status", regulatory.get("status", ""))),
        "regulatory_derogation_needed": str(digest.get("regulatory_derogation_needed", regulatory.get("derogation_needed", ""))),
        "detective_status": str(digest.get("detective_status", detective.get("status", ""))),
        "derogation_scan_overall": str(digest.get("derogation_scan_overall", derogation.get("overall", ""))),
        "verdict_per_anomaly_text": verdict_text,
        "citations_text": citations_text,
        "regulatory_findings_text": regulatory_findings_text,
        "document_generated_at": datetime.now(timezone.utc).isoformat(),
    }


def try_render_vdd_docx(*, repo_root: Path, final_state: dict[str, Any]) -> Path | None:
    """
    Render a Word VDD draft when enabled (Phase 4).

    - ``VDD_DOCX_PATH``: explicit output ``.docx`` path (relative to repo root or absolute).
    - Or ``VDD_DOCX=1`` to write ``output/vdd_pre_isa_draft.docx``.
    """
    path_raw = os.getenv("VDD_DOCX_PATH", "").strip()
    flag = os.getenv("VDD_DOCX", "0").strip().lower()

    if path_raw.lower() in ("-", "0", "false", "no", "off", "none"):
        return None

    if path_raw:
        out = Path(path_raw).expanduser()
        if not out.is_absolute():
            out = (repo_root / out).resolve()
    elif flag in {"1", "true", "yes", "on"}:
        out = (repo_root / "output" / "vdd_pre_isa_draft.docx").resolve()
    else:
        return None

    try:
        from docxtpl import DocxTemplate
    except ImportError:
        print("VDD_DOCX: docxtpl is not installed. Add it with: pip install docxtpl")
        return None

    template_path = ensure_default_vdd_template(repo_root)
    context = build_vdd_template_context(final_state)

    out.parent.mkdir(parents=True, exist_ok=True)
    tpl = DocxTemplate(str(template_path))
    tpl.render(context)
    tpl.save(str(out))

    if os.getenv("VDD_APPEND_REGULATORY_APPENDIX", "1").strip().lower() not in {
        "0",
        "false",
        "no",
        "off",
    }:
        _append_regulatory_appendix_to_docx(out, final_state)

    return out


def _append_regulatory_appendix_to_docx(out: Path, final_state: dict[str, Any]) -> None:
    """
    After Jinja render, append the regulatory PASS/FAIL list so Word always contains it
    without editing ``vdd_pre_isa_template.docx``. Disable with ``VDD_APPEND_REGULATORY_APPENDIX=0``.
    Context key ``regulatory_findings_text`` is still available for custom templates.
    """
    regulatory = final_state.get("regulatory_report") or {}
    if not isinstance(regulatory, dict):
        return
    rac = max(120, int(os.getenv("VDD_REGULATORY_RATIONALE_MAX_CHARS", "500") or "500"))
    body = format_regulatory_findings_plain(regulatory, rationale_max=rac).strip()
    if not body:
        return
    try:
        from docx import Document
    except ImportError:
        return
    doc = Document(str(out))
    doc.add_page_break()
    doc.add_heading("Appendix — CEI EN 50128 per-rule results", level=1)
    doc.add_paragraph(
        "Scoring uses keyword overlap between each rule (from cei_en_50128_rules.json) and evidence text "
        "from traceability anomalies, derogation scan, auditor, and detective. PASS = sufficient overlap; "
        "FAIL on MUST/SHALL = gap or formal derogation topic — not a semantic diff between two documents."
    )
    for block in body.split("\n\n"):
        if block.strip():
            doc.add_paragraph(block.replace("\n", " ").strip())
    doc.save(str(out))
