"""
Extract ``requirements_trace`` rows from an RSSOM-style Word document (e.g. FIT).

This is the same role as hand-editing ``data/requirements_trace.csv``, but filled from
the Verification Matrix + tables + body paragraphs. Output schema matches
:func:`core.project_ingestion.load_requirements_csv`.

Canonical file is ``data/requirements_trace.csv``. **Overwrite** from RSSOM only:

  python -m core.rssom_requirements_trace \\
    --docx data/RSSOM_APCS_FIT.docx \\
    --out data/requirements_trace.csv

**Merge** (enrich titles from RSSOM; keep manual-only rows and ``verification_status``):

  python -m core.rssom_requirements_trace \\
    --docx data/RSSOM_APCS_FIT.docx \\
    --out data/requirements_trace.csv \\
    --merge

Default requirement IDs: ``C6-<subsystem>-<number>`` (not only APCS), and
``C6_APCS_1`` → ``C6-APCS-1``. Override with ``--id-regex``.
"""
from __future__ import annotations

import argparse
import csv
import re
from collections import defaultdict
from collections.abc import Callable
from pathlib import Path
from typing import Any

from docx import Document

DEFAULT_REQUIREMENT_ID_REGEX = r"\b(?:C6-[A-Z0-9]+-\d+|C6_[A-Z0-9]+_\d+)\b"


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").strip())


def make_requirement_id_finder(pattern: str | None) -> Callable[[str], list[str]]:
    """Find requirement IDs in text; values uppercased, underscores → hyphens."""
    rx = re.compile(pattern or DEFAULT_REQUIREMENT_ID_REGEX, re.IGNORECASE)

    def find_ids(text: str) -> list[str]:
        out: list[str] = []
        for m in rx.finditer(text or ""):
            rid = m.group(0).upper().replace("_", "-")
            out.append(rid)
        return out

    return find_ids


def _clean_title(s: str) -> str:
    t = (
        s.replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u2018", "'")
        .replace("\u2019", "'")
    )
    t = t.replace('"', "")
    return _norm(t)


def _parse_verification_matrix(doc: Document, find_ids: Callable[[str], list[str]]) -> dict[str, list[str]]:
    rid_to_titles: dict[str, list[str]] = defaultdict(list)
    for table in doc.tables:
        if len(table.rows) < 2 or len(table.rows[0].cells) < 3:
            continue
        hdr = [_norm(c.text) for c in table.rows[0].cells[:3]]
        if not hdr[2].lower().startswith("requirements"):
            continue
        if "test case" not in hdr[0].lower():
            continue
        for row in table.rows[1:]:
            cells = [_norm(c.text) for c in row.cells]
            if len(cells) < 3:
                continue
            title = cells[1]
            req_cell = cells[2]
            ids = find_ids(req_cell)
            if not title or not ids:
                continue
            for rid_u in ids:
                if title not in rid_to_titles[rid_u]:
                    rid_to_titles[rid_u].append(title)
    return rid_to_titles


def _paragraph_contexts_by_requirement_id(
    doc: Document, find_ids: Callable[[str], list[str]]
) -> dict[str, list[str]]:
    rid_to_paras: dict[str, list[str]] = defaultdict(list)
    for para in doc.paragraphs:
        raw = _norm(para.text)
        if not raw:
            continue
        ids = find_ids(raw)
        if not ids:
            continue
        snippet = raw[:2000]
        for rid_u in ids:
            if snippet not in rid_to_paras[rid_u]:
                rid_to_paras[rid_u].append(snippet)
    return rid_to_paras


def _fallback_titles_from_doc(doc: Document, find_ids: Callable[[str], list[str]]) -> dict[str, str]:
    rid_to_desc: dict[str, str] = {}
    for table in doc.tables:
        for row in table.rows:
            cells = [_norm(c.text) for c in row.cells]
            line = " | ".join(cells)
            ids = find_ids(line)
            if not ids:
                continue
            for rid_u in ids:
                parts = [c for c in cells if c and rid_u not in c.upper() and len(c) > 12]
                desc = max(parts, key=len) if parts else line
                if rid_u not in rid_to_desc or len(desc) > len(rid_to_desc[rid_u]):
                    rid_to_desc[rid_u] = desc[:2000]
    return rid_to_desc


def rows_from_rssom_docx(
    docx_path: Path,
    *,
    id_regex: str | None = None,
    default_status: str = "PLANNED",
) -> list[dict[str, str]]:
    """Return rows ``{requirement_id, title, verification_status}`` from a FIT/RSSOM .docx."""
    if id_regex:
        try:
            re.compile(id_regex)
        except re.error as e:
            raise ValueError(f"Invalid id_regex: {e}") from e
    find_ids = make_requirement_id_finder(id_regex)
    doc = Document(str(docx_path))
    matrix = _parse_verification_matrix(doc, find_ids)
    para_ctx = _paragraph_contexts_by_requirement_id(doc, find_ids)
    fallback = _fallback_titles_from_doc(doc, find_ids)

    all_ids = sorted(
        set(matrix.keys()) | set(para_ctx.keys()) | set(fallback.keys()),
        key=_requirement_id_sort_key,
    )
    rows: list[dict[str, str]] = []
    for rid in all_ids:
        chunks: list[str] = []
        chunks.extend(matrix.get(rid, []))
        chunks.extend(para_ctx.get(rid, []))
        if rid in fallback:
            chunks.append(fallback[rid])
        seen: set[str] = set()
        merged: list[str] = []
        for c in chunks:
            cl = _clean_title(c)
            if not cl or cl in seen:
                continue
            seen.add(cl)
            merged.append(cl)
        title = " | ".join(merged)
        if len(title) > 500:
            title = title[:497] + "..."
        rows.append(
            {
                "requirement_id": rid,
                "title": title,
                "verification_status": default_status,
            }
        )
    return rows


def _requirement_id_sort_key(rid: str) -> tuple:
    m = re.search(r"(\d+)$", rid)
    tail = int(m.group(1)) if m else 0
    return (rid.split("-")[0] if "-" in rid else rid, tail, rid)


def merge_requirements_with_rssom(
    existing: list[dict[str, str]],
    rssom_rows: list[dict[str, str]],
) -> list[dict[str, str]]:
    """
    Enrich trace rows with RSSOM metadata without dropping manual-only IDs.

    - For each existing row whose ``requirement_id`` appears in ``rssom_rows``, **title** is
      replaced with the RSSOM-derived title; **verification_status** is kept from the CSV row.
    - Rows whose IDs never appear in RSSOM (e.g. synthetic eval IDs) are copied unchanged,
      including duplicate rows for the same ID.
    - IDs present only in RSSOM are appended (sorted) after all existing rows.
    """
    rssom_by_id = {str(r["requirement_id"]).strip().upper(): r for r in rssom_rows if r.get("requirement_id")}
    existing_ids = {
        str(r.get("requirement_id", "")).strip().upper()
        for r in existing
        if str(r.get("requirement_id", "")).strip()
    }

    out: list[dict[str, str]] = []
    for row in existing:
        rid = str(row.get("requirement_id", "")).strip().upper()
        if not rid:
            continue
        base = {k: str(v) if v is not None else "" for k, v in row.items()}
        if rid in rssom_by_id:
            base["title"] = rssom_by_id[rid]["title"]
        out.append(base)

    rssom_only = [
        {
            "requirement_id": r["requirement_id"],
            "title": r["title"],
            "verification_status": r.get("verification_status", "PLANNED"),
        }
        for r in rssom_rows
        if str(r.get("requirement_id", "")).strip().upper() not in existing_ids
    ]
    rssom_only.sort(key=lambda x: _requirement_id_sort_key(str(x["requirement_id"]).upper()))
    return out + rssom_only


def write_requirements_trace_csv(out_path: Path, rows: list[dict[str, Any]]) -> None:
    """Write UTF-8 CSV with columns requirement_id, title, verification_status."""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(["requirement_id", "title", "verification_status"])
        for r in rows:
            w.writerow(
                [
                    str(r.get("requirement_id", "")),
                    str(r.get("title", "")),
                    str(r.get("verification_status", "")),
                ]
            )


def _cli() -> None:
    from core.project_ingestion import load_requirements_csv

    ap = argparse.ArgumentParser(description="Build requirements_trace CSV from RSSOM/FIT .docx.")
    ap.add_argument("--docx", type=Path, required=True, help="Path to RSSOM_APCS_FIT.docx (or similar)")
    ap.add_argument(
        "--out",
        type=Path,
        required=True,
        help="Output CSV path (typically data/requirements_trace.csv — overwrites).",
    )
    ap.add_argument("--default-status", default="PLANNED", help="verification_status for rows taken only from RSSOM")
    ap.add_argument("--id-regex", default=None, metavar="PATTERN", help="Override requirement ID regex")
    ap.add_argument(
        "--merge",
        action="store_true",
        help="If --out exists, merge: update titles from RSSOM for matching IDs, keep statuses & manual-only rows; append new RSSOM-only IDs.",
    )
    args = ap.parse_args()
    rssom_rows = rows_from_rssom_docx(
        args.docx,
        id_regex=args.id_regex,
        default_status=args.default_status,
    )
    if args.merge:
        existing: list[dict[str, str]] = []
        if args.out.is_file():
            existing = load_requirements_csv(args.out)
        rows = merge_requirements_with_rssom(existing, rssom_rows)
    else:
        rows = rssom_rows
    write_requirements_trace_csv(args.out, rows)
    print(f"Wrote {len(rows)} rows to {args.out}")


if __name__ == "__main__":
    _cli()
